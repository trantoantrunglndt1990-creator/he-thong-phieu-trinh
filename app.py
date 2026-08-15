import streamlit as st
import google.generativeai as genai
from docxtpl import DocxTemplate
import pandas as pd
import io
import json
from datetime import datetime
import docx

# Cấu hình giao diện
st.set_page_config(page_title="Hệ thống Sản xuất Phiếu trình", layout="wide")
st.title("📄 Trợ lý AI - Tự động hóa Phiếu trình & Danh mục công việc")

# 1. Cấu hình API an toàn
api_key = st.secrets.get("GEMINI_API_KEY", "")
if not api_key:
    st.error("⚠️ Lỗi: Chưa cấu hình GEMINI_API_KEY. Vui lòng kiểm tra lại cài đặt Secrets trên Streamlit Cloud.")
    st.stop()

# Khởi tạo mô hình
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash")

# 2. Khởi tạo Bảng Danh mục công việc lưu trữ tạm thời
if 'danh_muc' not in st.session_state:
    st.session_state.danh_muc = pd.DataFrame(columns=["Ngày", "Số Phiếu", "Nội dung công việc", "Thời hạn", "Trạng thái"])

col1, col2 = st.columns([1, 1])

with col1:
    st.header("1. Tải tài liệu nguồn")
    uploaded_file = st.file_uploader("Tải công văn, tờ trình... (DOCX, PDF, TXT)", type=["txt", "pdf", "docx"])

    if uploaded_file and st.button("🚀 Phân tích & Trích xuất"):
        with st.spinner("Hệ thống đang xử lý tài liệu. Vui lòng đợi..."):
            try:
                # Chuẩn bị câu lệnh Prompt
                prompt = """
                Đóng vai một chuyên viên hành chính mẫn cán. Đọc tài liệu đính kèm và trích xuất thông tin:
                1. noi_dung_trinh: Tên sự việc chính cần trình Lãnh đạo.
                2. can_cu_phap_ly: Liệt kê số hiệu, ngày, tên cơ quan của các văn bản liên quan.
                3. tom_tat_nhiem_vu: Tóm tắt ngắn gọn yêu cầu, nhiệm vụ (1-2 câu).
                4. thoi_han: Thời hạn hoàn thành (nếu có, không có ghi "Không quy định").
                5. de_xuat_giai_quyet: Tên văn bản đề xuất Lãnh đạo ký.
                Trả về DUY NHẤT một chuỗi JSON hợp lệ, không kèm giải thích hay ký tự thừa: 
                {"noi_dung_trinh": "...", "can_cu_phap_ly": "...", "tom_tat_nhiem_vu": "...", "thoi_han": "...", "de_xuat_giai_quyet": "..."}
                """

                # Biến lưu trữ dữ liệu truyền cho AI
                contents = []

                # Xử lý theo từng định dạng file
                if uploaded_file.name.endswith(".docx"):
                    doc_reader = docx.Document(uploaded_file)
                    text_data = [para.text for para in doc_reader.paragraphs if para.text.strip()]
                    # Quét thêm chữ nằm trong các bảng biểu
                    for table in doc_reader.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_data.append(cell.text.strip())
                    
                    file_content = "\n".join(text_data)
                    if not file_content.strip():
                        st.warning("⚠️ File Word trống hoặc không chứa văn bản có thể đọc.")
                        st.stop()
                    contents = [file_content, prompt]

                elif uploaded_file.name.endswith(".txt"):
                    file_content = uploaded_file.read().decode("utf-8")
                    contents = [file_content, prompt]

                else: # Mặc định là PDF
                    file_bytes = uploaded_file.read()
                    contents = [{"mime_type": "application/pdf", "data": file_bytes}, prompt]

                # Gọi API của Google Gemini
                response = model.generate_content(contents)
                
                # Làm sạch và chuyển đổi kết quả JSON
                raw_text = response.text.replace("```json", "").replace("```", "").strip()
                st.session_state.temp_data = json.loads(raw_text)
                st.success("✅ Phân tích thành công! Mời rà soát kết quả ở bảng bên cạnh.")

            except Exception as e:
                # Bắt mọi lỗi và hiển thị thông báo tiếng Việt
                st.error(f"❌ Có lỗi xảy ra trong quá trình AI phân tích.")
                st.error(f"Chi tiết lỗi kỹ thuật: {str(e)}")
                st.info("💡 Gợi ý: Máy chủ Google có thể đang bận hoặc file tải lên bị lỗi cấu trúc. Hãy thử bấm phân tích lại.")

with col2:
    st.header("2. Hoàn thiện Phiếu trình")
    if 'temp_data' in st.session_state:
        # Giao diện nhập liệu chỉnh sửa
        so_phieu = st.text_input("Số Phiếu trình:", value="")
        noi_dung = st.text_area("Nội dung trình:", value=st.session_state.temp_data.get("noi_dung_trinh", ""))
        can_cu = st.text_area("Căn cứ pháp lý:", value=st.session_state.temp_data.get("can_cu_phap_ly", ""))
        tom_tat = st.text_area("Tóm tắt nhiệm vụ:", value=st.session_state.temp_data.get("tom_tat_nhiem_vu", ""))
        thoi_han = st.text_input("Thời hạn xử lý:", value=st.session_state.temp_data.get("thoi_han", ""))
        de_xuat = st.text_area("Đề xuất giải quyết:", value=st.session_state.temp_data.get("de_xuat_giai_quyet", ""))

        if st.button("💾 Xuất Phiếu Trình & Lưu Báo Cáo"):
            try:
                now = datetime.now()
                
                # Điền dữ liệu vào file Word template.docx
                doc = DocxTemplate("template.docx")
                doc.render({
                    "so_phieu": so_phieu,
                    "ngay": now.strftime("%d"),
                    "thang": now.strftime("%m"),
                    "nam": now.strftime("%Y"),
                    "noi_dung_trinh": noi_dung,
                    "can_cu_phap_ly": can_cu,
                    "tom_tat_nhiem_vu": tom_tat,
                    "thoi_han": thoi_han,
                    "de_xuat_giai_quyet": de_xuat
                })
                
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # Cập nhật thêm dòng mới vào danh mục
                new_row = {
                    "Ngày": now.strftime("%d/%m/%Y"), 
                    "Số Phiếu": so_phieu, 
                    "Nội dung công việc": noi_dung, 
                    "Thời hạn": thoi_han, 
                    "Trạng thái": "Đang trình ký"
                }
                st.session_state.danh_muc = pd.concat([st.session_state.danh_muc, pd.DataFrame([new_row])], ignore_index=True)
                
                st.success("✅ Đã tạo Phiếu trình và lưu danh mục thành công!")
                st.download_button("📥 Tải Phiếu Trình (.docx)", data=doc_io, file_name=f"Phieu_Trinh_{so_phieu}.docx")
            
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo file Word: {str(e)}")
                st.info("💡 Gợi ý: Hãy kiểm tra chắc chắn bạn đã tải file 'template.docx' lên hệ thống GitHub chưa.")

st.divider()
st.header("3. Danh mục công việc thực hiện (Phục vụ báo cáo tuần/tháng)")
st.dataframe(st.session_state.danh_muc, use_container_width=True)

if not st.session_state.danh_muc.empty:
    excel_io = io.BytesIO()
    st.session_state.danh_muc.to_excel(excel_io, index=False)
    excel_io.seek(0)
    st.download_button("📊 Tải file Danh mục (Excel)", data=excel_io, file_name="Danh_Muc_Cong_Viec.xlsx")
